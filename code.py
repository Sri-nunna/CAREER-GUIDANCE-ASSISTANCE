import streamlit as st
import google.generativeai as genai
import speech_recognition as sr
from gtts import gTTS
import os
import tempfile
import base64
from io import BytesIO
from PIL import Image, ImageDraw, ImageFont
import pandas as pd
import plotly.express as px
from docx import Document

# Set page config as the first command
st.set_page_config(page_title="AI Career Counselor", layout="wide", page_icon="💼")

# Custom CSS for a single background image across all tabs
st.markdown(
    """
    <style>
    /* General background and font */
    [data-testid="stAppViewContainer"] {
        font-family: 'Arial', sans-serif;
        background-image: url('https://humanrights.gov.au/sites/default/files/styles/full_content_hero_image/public/2023-08/adobestock_422946225.jpeg?itok=rMLAU6_6');
        background-size: cover;
        background-repeat: no-repeat;
        background-attachment: fixed;
        background-position: center;
    }

    /* Ensure the Streamlit app content is visible */
    .main .block-container {
        background-color: rgba(255, 255, 255, 0.8); /* Semi-transparent white background */
        padding: 2rem;
        border-radius: 10px;
    }

    /* Sidebar styling */
    .sidebar .sidebar-content {
        background: rgba(44, 62, 80, 0.9); /* Add transparency to sidebar */
        color: white;
        padding: 20px;
        border-radius: 10px;
    }

    /* Button styling */
    .stButton>button {
        background: #3498db;
        color: white;
        border: none;
        border-radius: 5px;
        padding: 10px 20px;
        font-size: 16px;
        transition: background 0.3s ease;
    }

    .stButton>button:hover {
        background: #2980b9;
    }

    /* Input field styling - Remove borders and white grid */
    .stTextInput>div>div>input, .stTextArea>div>div>textarea {
        background: transparent;
        border: none; /* Remove border */
        border-radius: 0; /* Remove rounded corners */
        padding: 10px;
        width: 100%;
        color: white; /* Ensure text is visible */
        box-shadow: none; /* Remove any shadow */
    }

    .stTextInput>div>div>input::placeholder, .stTextArea>div>div>textarea::placeholder {
        color: #666; /* Placeholder text color */
    }

    /* File uploader styling */
    .stFileUploader>div>div>div {
        background: transparent;
        border: none; /* Remove border */
        border-radius: 0; /* Remove rounded corners */
        padding: 10px;
        box-shadow: none; /* Remove any shadow */
    }

    /* Success and warning messages */
    .stSuccess {
        background: rgba(212, 237, 218, 0.9); /* Add transparency */
        color: #155724;
        border-radius: 5px;
        padding: 10px;
    }

    .stWarning {
        background: rgba(255, 243, 205, 0.9); /* Add transparency */
        color: #856404;
        border-radius: 5px;
        padding: 10px;
    }

    /* Header styling */
    h1 {
        font-size: 36px !important;
        color: #2c3e50;
        margin-bottom: 20px;
    }

    h2 {
        font-size: 28px !important;
        color: #2c3e50;
        margin-bottom: 15px;
    }

    h3 {
        font-size: 22px !important;
        color: #2c3e50;
        margin-bottom: 10px;
    }

    /* Spacing for better readability */
    .stTabs>div>div>div>div {
        font-size: 18px !important;
    }

    .stRadio>div>div {
        margin-bottom: 20px;
    }

    .stTextInput>div>div>input, .stTextArea>div>div>textarea {
        margin-bottom: 15px;
    }

    .stFileUploader>div>div>div {
        margin-bottom: 15px;
    }
    </style>
    """,
    unsafe_allow_html=True,
)

# Configure Google Gemini AI Key
API_KEY = "AIzaSyBXbeGn8nLtLnjAEbDlzrmIIATXSuycX44"  # Replace with your actual API key
if API_KEY == "YOUR_GOOGLE_GEMINI_API_KEY":
    st.sidebar.error("⚠ Please set a valid Google Gemini API key!")
else:
    genai.configure(api_key=API_KEY)
    model = genai.GenerativeModel("gemini-2.0-flash")


# Function: Convert Text to Speech
def speak(text, lang="en"):
    if text.strip():
        tts = gTTS(text=text, lang=lang)
        with tempfile.NamedTemporaryFile(delete=False, suffix=".mp3") as tmpfile:
            audio_path = tmpfile.name
            tts.save(audio_path)
        st.audio(audio_path, format="audio/mp3")
        os.unlink(audio_path)


# Function: Convert Speech to Text
def get_speech_input(language="en"):
    recognizer = sr.Recognizer()
    with sr.Microphone() as source:
        st.info("🎤 Speak now...")
        recognizer.adjust_for_ambient_noise(source)
        try:
            audio = recognizer.listen(source, timeout=5, phrase_time_limit=5)
            text = recognizer.recognize_google(audio, language=language)
            return text
        except sr.UnknownValueError:
            return "⚠ Could not understand the audio."
        except sr.RequestError:
            return "⚠ Speech recognition service is unavailable."


# Function: Process Image for AI Analysis
def process_image(image):
    img = Image.open(image).convert("RGB")
    buffered = BytesIO()
    img.save(buffered, format="PNG")
    return base64.b64encode(buffered.getvalue()).decode()


# Function: Extract text from resume
def extract_resume_text(file):
    if file.type == "application/pdf":
        return file.read().decode("utf-8", errors="ignore")
    elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        doc = Document(file)
        return "\n".join([para.text for para in doc.paragraphs])
    return file.read().decode("utf-8", errors="ignore")


# Function: Generate Professional Resume as an Image
def generate_resume_image(name, contact_info, skills, experience, education, projects, certifications, languages):
    # Create a blank image with white background
    width, height = 800, 1400  # Increased height for additional sections
    resume_image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(resume_image)

    # Load fonts (use a custom font if available)
    try:
        title_font = ImageFont.truetype("arial.ttf", 30)
        header_font = ImageFont.truetype("arial.ttf", 24)
        body_font = ImageFont.truetype("arial.ttf", 18)
    except IOError:
        title_font = ImageFont.load_default()
        header_font = ImageFont.load_default()
        body_font = ImageFont.load_default()

    # Define colors
    header_color = (70, 130, 180)  # Steel blue
    section_color = (220, 220, 220)  # Light gray
    accent_color = (0, 128, 128)  # Teal

    # Add a header with a background color
    draw.rectangle([(0, 0), (width, 100)], fill=header_color)
    draw.text((50, 30), "RESUME", font=title_font, fill="white")

    # Add Name and Contact Info
    draw.text((50, 120), f"Name: {name}", font=header_font, fill="black")
    draw.text((50, 160), f"Contact: {contact_info}", font=body_font, fill="black")

    # Add Skills Section
    draw.rectangle([(50, 220), (width - 50, 240)], fill=section_color)
    draw.text((50, 220), "Skills", font=header_font, fill=accent_color)
    draw.text((50, 260), skills, font=body_font, fill="black")

    # Add Experience Section
    draw.rectangle([(50, 340), (width - 50, 360)], fill=section_color)
    draw.text((50, 340), "Experience", font=header_font, fill=accent_color)
    draw.text((50, 380), experience, font=body_font, fill="black")

    # Add Education Section
    draw.rectangle([(50, 500), (width - 50, 520)], fill=section_color)
    draw.text((50, 500), "Education", font=header_font, fill=accent_color)
    draw.text((50, 540), education, font=body_font, fill="black")

    # Add Projects Section
    draw.rectangle([(50, 620), (width - 50, 640)], fill=section_color)
    draw.text((50, 620), "Projects", font=header_font, fill=accent_color)
    draw.text((50, 660), projects, font=body_font, fill="black")

    # Add Certifications Section
    draw.rectangle([(50, 740), (width - 50, 760)], fill=section_color)
    draw.text((50, 740), "Certifications", font=header_font, fill=accent_color)
    draw.text((50, 780), certifications, font=body_font, fill="black")

    # Add Languages Section
    draw.rectangle([(50, 860), (width - 50, 880)], fill=section_color)
    draw.text((50, 860), "Languages", font=header_font, fill=accent_color)
    draw.text((50, 900), languages, font=body_font, fill="black")

    # Add a footer
    footer_color = (70, 130, 180)  # Steel blue
    draw.rectangle([(0, height - 50), (width, height)], fill=footer_color)
    draw.text((50, height - 40), "Generated by AI Career Counselor", font=body_font, fill="white")

    # Save the image to a BytesIO object
    buffered = BytesIO()
    resume_image.save(buffered, format="PNG")
    return buffered.getvalue()


# UI Design Enhancements
st.sidebar.title("💼 AI Career Counselor")
st.sidebar.write("🚀 Get career guidance, resume analysis, and job recommendations!")

# 🌍 Features List in Sidebar
st.sidebar.header("✨ Features")
st.sidebar.write("""
- **🏆 Career Guidance**
- **📄 Resume Analyzer**
- **📊 Job Market Insights**
- **🎤 Voice Query**
- **📸 Image Analysis**
- **📚 Learning Path**
- **🎤 Mock Interview**
- **🤝 Networking**
""")

# 🔹 Main Tabs
tabs = st.tabs(["🏆 Career Guidance", "📄 Resume Analyzer", "📊 Job Market Insights", "🎤 Voice Query", "📸 Image Analysis",
                "📚 Learning Path", "🎤 Mock Interview", "🤝 Networking"])

# 🏆 Career Guidance Tab
with tabs[0]:
    st.header("🏆 AI-Powered Career Guidance")
    user_query = st.text_input("🔍 Enter your career-related question:")
    if st.button("🤖 Get AI Advice"):
        if user_query:
            with st.spinner("⏳ AI is analyzing..."):
                response = model.generate_content([user_query])
                st.success("✅ AI Response:")
                st.write(response.text)
                speak(response.text, lang="en")  # Default language for text-to-speech
        else:
            st.warning("⚠ Please enter a question!")

# 📄 Resume Analyzer Tab
with tabs[1]:
    st.header("📄 Resume Analyzer & Builder")
    resume_option = st.radio("Choose an option:", ["📄 Analyze Existing Resume", "🛠 Build & Optimize Resume"])

    if resume_option == "📄 Analyze Existing Resume":
        uploaded_resume = st.file_uploader("🔼 Upload your resume (PDF/DOCX/TXT)", type=["pdf", "docx", "txt"])
        if uploaded_resume:
            st.info("📊 Analyzing resume...")
            with st.spinner("🧠 AI is evaluating your resume..."):
                resume_text = extract_resume_text(uploaded_resume)
                response = model.generate_content(["Analyze and improve this resume:", resume_text])
                st.success("✅ Resume Feedback:")
                st.write(response.text)
                speak(response.text, lang="en")  # Default language for text-to-speech

    elif resume_option == "🛠 Build & Optimize Resume":
        st.subheader("🛠 Build & Optimize Your Resume")
        name = st.text_input("👤 Full Name")
        contact_info = st.text_input("📞 Contact Info (e.g., Phone, Email)")
        skills = st.text_area("🛠 Skills (comma-separated)")
        experience = st.text_area("💼 Experience (e.g., Job Title, Company, Duration, Responsibilities)")
        education = st.text_area("🎓 Education (e.g., Degree, Institution, Year)")
        projects = st.text_area("📂 Projects (e.g., Project Title, Description, Technologies Used)")
        certifications = st.text_area("📜 Certifications (e.g., Certification Name, Issuing Organization)")
        languages = st.text_area("🌍 Languages (e.g., English, Spanish)")

        if st.button("🚀 Generate Resume"):
            if name and contact_info and skills and experience and education and projects and languages:
                with st.spinner("⏳ Generating your resume..."):
                    resume_image = generate_resume_image(name, contact_info, skills, experience, education, projects,
                                                         certifications, languages)
                    st.success("✅ Your Resume is Ready!")
                    st.image(resume_image, caption="Generated Resume", use_container_width=True)

                    # Add a download button for the resume
                    st.download_button(
                        label="📥 Download Resume",
                        data=resume_image,
                        file_name="resume.png",
                        mime="image/png"
                    )
            else:
                st.warning("⚠ Please fill in all fields!")

# 📊 Job Market Insights Tab
with tabs[2]:
    st.header("📊 Real-Time Job Market Insights")
    job_data = pd.DataFrame({
        "Industry": ["Technology", "Healthcare", "Finance", "Education", "Retail"],
        "Vacancies": [15000, 12000, 8000, 5000, 7000]
    })
    fig = px.bar(job_data, x="Industry", y="Vacancies", title="Job Vacancies by Industry", color="Industry")
    st.plotly_chart(fig)

    st.subheader("💡 Skill Recommendations")
    selected_industry = st.selectbox("Choose an industry:", job_data["Industry"].tolist())
    if st.button("🔍 Get Skill Recommendations"):
        with st.spinner("🔎 Analyzing trends..."):
            skill_response = model.generate_content(f"Suggest top skills needed for {selected_industry} industry.")
            st.success("✅ Recommended Skills:")
            st.write(skill_response.text)
            speak(skill_response.text, lang="en")  # Default language for text-to-speech

# 🎤 Voice Query Tab
with tabs[3]:
    st.header("🎤 Ask a Career Question Using Voice")

    # Language selection for voice query
    voice_language = st.selectbox("🌐 Select a language for voice query:",
                                  ["en", "te", "hi", "es", "fr", "de", "zh", "ja"])

    if st.button("🎙 Start Speaking"):
        speech_text = get_speech_input(language=voice_language)
        st.write(f"🗣 You said: **{speech_text}**")
        if speech_text:
            with st.spinner("🤔 AI is analyzing your query..."):
                response = model.generate_content([speech_text])
                st.success("✅ AI Response:")
                st.write(response.text)
                speak(response.text, lang=voice_language)  # Convert AI response to speech in the selected language

# 📸 Image Analysis Tab
with tabs[4]:
    st.header("📸 Upload an Image for AI Analysis")
    uploaded_image = st.file_uploader("🔼 Upload an image (e.g., job role charts, skill maps)",
                                      type=["jpg", "png", "jpeg"])
    if uploaded_image:
        st.image(uploaded_image, caption="🖼 Uploaded Image", use_container_width=True)
        st.info("🔍 Processing image...")
        with st.spinner("📡 AI is analyzing your image..."):
            image_base64 = process_image(uploaded_image)
            response = model.generate_content(
                ["Describe this image:", {"mime_type": "image/png", "data": image_base64}])
            st.success("✅ AI Response:")
            st.write(response.text)
            speak(response.text, lang="en")  # Default language for text-to-speech

# 📚 Learning Path Tab
with tabs[5]:
    st.header("📚 Personalized Learning Path & Skill Recommendations")
    user_skills = st.text_area("🛠 Enter your current skills (comma-separated):")
    target_role = st.text_input("🎯 Enter your target job role:")
    if st.button("📚 Get Learning Path"):
        if user_skills and target_role:
            with st.spinner("🔍 Analyzing skill gaps..."):
                response = model.generate_content(
                    [f"Suggest a learning path for someone with skills: {user_skills} to become a {target_role}."])
                st.success("✅ Learning Path Recommendations:")
                st.write(response.text)
                speak(response.text, lang="en")  # Default language for text-to-speech
        else:
            st.warning("⚠ Please enter your skills and target role!")

# 🎤 Mock Interview Tab
with tabs[6]:
    st.header("🎤 AI-Powered Mock Interview & Personality Analysis")
    interview_role = st.text_input("🎯 Enter the job role for the mock interview:")

    if interview_role:
        st.success(f"🎯 Mock Interview for **{interview_role}** Role")

        # Generate dynamic questions based on the job role
        with st.spinner("🤖 Generating interview questions..."):
            questions_prompt = f"Generate 30 interview questions for a {interview_role} role. Each question should have 4 options (A, B, C, D) and a correct answer. Format each question as follows:\n\nQuestion: [Question text]\nA) [Option A]\nB) [Option B]\nC) [Option C]\nD) [Option D]\nCorrect Answer: [Correct Option]"
            questions_response = model.generate_content(questions_prompt)
            questions_text = questions_response.text

        # Parse the generated questions into a structured format
        questions = []
        question_blocks = questions_text.split("\n\n")  # Split by double newlines
        for block in question_blocks:
            lines = block.split("\n")
            if len(lines) >= 6:  # Ensure the block has a question, 4 options, and a correct answer
                question = lines[0].replace("Question: ", "").strip()
                options = [line.strip() for line in lines[1:5]]
                correct_answer = lines[5].replace("Correct Answer: ", "").strip()
                questions.append({
                    "question": question,
                    "options": options,
                    "correct_answer": correct_answer
                })

        # Generate a document with the questions and answers
        if st.button("📄 Download Mock Interview Questions"):
            doc = Document()
            doc.add_heading(f"Mock Interview Questions for {interview_role}", level=1)
            doc.add_paragraph("Below are the questions along with their correct answers:")

            for i, q in enumerate(questions):
                doc.add_paragraph(f"Question {i + 1}: {q['question']}")
                for option in q["options"]:
                    doc.add_paragraph(option)
                doc.add_paragraph(f"Correct Answer: {q['correct_answer']}")
                doc.add_paragraph()  # Add a blank line between questions

            # Save the document to a BytesIO object
            doc_bytes = BytesIO()
            doc.save(doc_bytes)
            doc_bytes.seek(0)

            # Provide a download button for the document
            st.download_button(
                label="📥 Download Mock Interview Questions",
                data=doc_bytes,
                file_name=f"Mock_Interview_{interview_role}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
# 🤝 Networking Tab
with tabs[7]:
    st.header("🤝 AI-Powered Professional Networking & Mentorship Suggestions")
    user_industry = st.text_input("🏭 Enter your industry:")
    if st.button("🔍 Get Networking Suggestions"):
        if user_industry:
            with st.spinner("🔍 Analyzing networking opportunities..."):
                response = model.generate_content([
                                                      f"Suggest professional networking and mentorship opportunities for someone in the {user_industry} industry."])
                st.success("✅ Networking Suggestions:")
                st.write(response.text)
                speak(response.text, lang="en")  # Default language for text-to-speech
        else:
            st.warning("⚠ Please enter your industry!")

st.success("🚀 AI Career Counselor is ready to assist you!")
